import pandas as pd
from anytree import Node, RenderTree
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def create_tree(data):
    root = Node("Root")
    nodes = {"Root": root}

    for row in data:
        for i, item in enumerate(row):
            if pd.notna(item):  # 检查项目是否不是 NaN
                parent = nodes[row[i - 1]] if i > 0 and pd.notna(row[i - 1]) else root
                if item not in nodes:
                    nodes[item] = Node(str(item), parent=parent)  # 确保 item 是字符串

    return root


def build_nested_list(node):
    if not node.children:
        return node.name
    return [node.name, [build_nested_list(child) for child in node.children]]


def add_to_document(document, item, level):
    if isinstance(item, str):
        document.add_paragraph(item, style=f'Heading {min(level, 9)}')
    elif isinstance(item, (list, tuple)) and len(item) > 0:
        document.add_paragraph(str(item[0]), style=f'Heading {min(level, 9)}')
        if len(item) > 1 and isinstance(item[1], (list, tuple)):
            for subitem in item[1]:
                add_to_document(document, subitem, level + 1)


# 读取Excel文件
df = pd.read_excel('初设模板标题.xlsx', header=None)  # 使用你的Excel文件名
data = df.values.tolist()

# 创建树结构
root = create_tree(data)

# 创建嵌套列表
nested_list = build_nested_list(root)
nested_list = nested_list[1]  # 移除顶层的 "Root" 节点

# 创建Word文档
doc = Document()

# 修改现有的标题样式
for i in range(1, 10):  # 为1级到9级标题修改样式
    style = doc.styles[f'Heading {i}']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(max(16 - i, 8))  # 确保字体大小不小于8pt
    font.bold = True

# 添加内容到文档
for item in nested_list:
    add_to_document(doc, item, 1)

# 保存文档
doc.save('output.docx')

print("Word文档已生成：output.docx")