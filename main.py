from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建一个新的Word文档
doc = Document()

# 定义“标书标题一”样式并应用自定义设置
style_title1 = doc.styles.add_style('标书标题一', WD_STYLE_TYPE.PARAGRAPH)
style_title1.base_style = doc.styles['Heading 1']
style_title1.font.name = 'Calibri'
style_title1.font.size = Pt(16)
style_title1.font.bold = True
style_title1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 定义“标题二”样式并应用自定义设置
style_title2 = doc.styles.add_style('标书标题二', WD_STYLE_TYPE.PARAGRAPH)
style_title2.base_style = doc.styles['Heading 2']
style_title2.font.name = 'Calibri'
style_title2.font.size = Pt(14)
style_title2.font.bold = True
style_title2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 定义“标题三”样式并应用自定义设置
style_title3 = doc.styles.add_style('标书标题三', WD_STYLE_TYPE.PARAGRAPH)
style_title3.base_style = doc.styles['Heading 3']
style_title3.font.name = 'Calibri'
style_title3.font.size = Pt(12)
style_title3.font.bold = True
style_title3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 定义“标题四”样式并应用自定义设置
style_title4 = doc.styles.add_style('标书标题四', WD_STYLE_TYPE.PARAGRAPH)
style_title4.base_style = doc.styles['Heading 4']
style_title4.font.name = 'Calibri'
style_title4.font.size = Pt(11)
style_title4.font.bold = True
style_title4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT



# 定义“标书正文”样式并应用自定义设置
style_body = doc.styles.add_style('标书正文', WD_STYLE_TYPE.PARAGRAPH)
style_body.font.name = 'Times New Roman'  # 通常用于正式文档
style_body.font.size = Pt(12)  # 常见的正文字体大小
style_body.font.bold = False  # 正文通常不加粗
style_body.paragraph_format.line_spacing = 1.5  # 设置1.5倍行距
style_body.paragraph_format.space_after = Pt(8)  # 段后添加8磅的间距
style_body.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐

# 使用“标书正文”样式添加一段示例正文
# 使用定义的样式添加示例标题
doc.add_paragraph('这是用“标书标题一”样式的标题', style='标书标题一')
doc.add_paragraph('这是用“标书标题二”样式的标题', style='标书标题二')
doc.add_paragraph('这是用“标书标题三”样式的标题', style='标书标题三')
doc.add_paragraph('这是用“标书标题四”样式的标题', style='标书标题四')
doc.add_paragraph('这是一段使用“标书正文”样式的文本。', style='标书正文')

#第一页插入目录引用


# 保存文档
doc.save('styled_document.docx')
