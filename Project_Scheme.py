from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 创建一个新的Word文档
doc = Document()

# 设计说明部分
design_section = doc.add_heading('设计说明', level=1)

# 建筑部分
building_paragraph = doc.add_paragraph()
building_run = building_paragraph.add_run('建筑')
building_run.bold = True

# 设计依据
design_basis_paragraph = doc.add_paragraph()
design_basis_run = design_basis_paragraph.add_run('设计依据')
design_basis_run.bold = True
design_basis_paragraph.level = 2

# 工程概况
project_summary_paragraph = doc.add_paragraph()
project_summary_run = project_summary_paragraph.add_run('工程概况')
project_summary_run.bold = True
project_summary_paragraph.level = 2

# 方案回顾
review_paragraph = doc.add_paragraph()
review_run = review_paragraph.add_run('方案回顾')
review_run.bold = True
review_paragraph.level = 2

# 面积明细表
area_detail_table_paragraph = doc.add_paragraph()
area_detail_table_run = area_detail_table_paragraph.add_run('面积明细表')
area_detail_table_run.bold = True
area_detail_table_paragraph.level = 2

# 建筑用料（单体）
material_paragraph = doc.add_paragraph()
material_run = material_paragraph.add_run('建筑用料（单体）')
material_run.bold = True
material_paragraph.level = 2

# 墙体
wall_paragraph = doc.add_paragraph()
wall_run = wall_paragraph.add_run('墙体')
wall_run.bold = True
wall_paragraph.level = 3

# 以此类推，为其他项目添加段落和适当的层级样式

# 保存文档
doc.save('Design Specifications.docx')