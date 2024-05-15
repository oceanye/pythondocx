import glob
import os


def create_folders_from_markdown_file(md_file_path, base_path):
    with open(md_file_path, 'r', encoding='utf-8') as file:
        markdown_text = file.read()

    lines = markdown_text.split('\n')
    current_path = base_path
    level_stack = []  # 保存各级标题对应的文件夹路径

    for line in lines:
        line = line.strip()  # 去除空白字符（包括空行）
        if not line:
            continue  # 如果是空行，则跳过

        level = len(line) - len(line.lstrip('#'))
        folder_name = line.lstrip('# ').strip()

        # 当前标题的文件夹路径
        current_folder_path = os.path.join(current_path, folder_name)

        # 将当前标题的文件夹路径放入正确的层级位置
        if level == len(level_stack):
            # 与栈顶层级相同，说明在同一层级
            pass
        elif level > len(level_stack):
            # 比栈顶层级高，进入下一级
            level_stack.append(current_path)
        else:
            # 比栈顶层级低，回到对应层级
            level_stack = level_stack[:level]
        # 更新当前路径
        current_path = os.path.join(*level_stack, folder_name)
        # 创建文件夹
        os.makedirs(current_path, exist_ok=True)


# Markdown 文件路径
md_file_path = '11.md'

# 指定路径
base_path = 'F:/test'

# 创建文件夹
create_folders_from_markdown_file(md_file_path, base_path)
# 获取文件夹路径下的所有文件名组成的列表
j=0
#path = 'F:/test'
# for root, dirs, files in os.walk('F:/test'):
#     for i in dirs:
#         print(root+i)
#         j+=1
# print(f'文件个数：{j}')


def build_keyword_to_folder_mapping(base_path):
    keyword_to_folder = {}
    # 遍历base_path下的所有文件夹
    for root, dirs, files in os.walk(base_path, topdown=True):
        for name in dirs:
            # 构建完整的路径
            folder_path = os.path.join(root, name)
            # 将文件夹名称作为关键字，映射到其完整路径
            keyword_to_folder[root] = folder_path
    return keyword_to_folder

# 构建映射字典
keyword_to_folder = build_keyword_to_folder_mapping(base_path)

# 打印出构建的字典以查看结果
print("Keyword to Folder Mapping:")
for keyword, folder_path in keyword_to_folder.items():
    print(f"Keyword: {keyword}, Folder: {folder_path}")

# 现在可以使用这个字典来处理文档内容
# 实现按标题建立文件夹，且拆分文档段落

import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def sanitize_filename(filename):
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename

def save_document(text, path, filename):
    filename = sanitize_filename(filename)
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.first_line_indent = Pt(24)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    text = text.lstrip()
    doc.add_paragraph(text, style=style)
    os.makedirs(path, exist_ok=True)
    full_path = os.path.join(path, f"{filename}.docx")
    try:
        doc.save(full_path)
    except Exception as e:
        print(f"Error saving document {full_path}: {e}")

def process_word_document(file_path):
    doc = Document(file_path)
    current_path = []
    base_path = "output"
    content_buffer = ""
    last_heading = ""

    for para in doc.paragraphs:
        style = para.style.name
        if 'Heading' in style:
            if content_buffer:
                path = os.path.join(base_path, *current_path)
                filename = last_heading[:50]  # Use first 50 characters of last heading as filename
                save_document(content_buffer, path, filename)
                content_buffer = ""
            level = int(style.split()[-1])
            current_path = current_path[:level-1]
            last_heading = sanitize_filename(para.text.strip())
            current_path.append(last_heading)
        elif 'Body Text' in style or 'Normal' in style:
            content_buffer += para.text + "\n"

    if content_buffer:
        path = os.path.join(base_path, *current_path)
        filename = last_heading[:50]  # Use first 50 characters of last heading as filename
        save_document(content_buffer, path, filename)

if __name__ == "__main__":
    process_word_document("通用扩初模板.docx")






























