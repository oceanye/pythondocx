# 导入必要的模块
from docx import Document

# 创建一个新的文档
document = Document()

# 添加文档标题
document.add_heading('Document Title', 0)

# 添加一个包含不同格式文本的段落
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# 添加二级标题
document.add_heading('Heading, level 1', level=1)

# 添加一个具有特定样式的引用段落
document.add_paragraph('Intense quote', style='Intense Quote')

# 添加无序列表和有序列表项
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# 注释掉添加图片的代码行
#document.add_picture('monty-truth.png', width=Inches(1.25))

# 定义数据记录，用于创建表格
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# 添加一个表格，并设置表头
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

# 循环，填充表格的数据行
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

# 添加一个分页符
document.add_page_break()

# 保存文档
document.save('demo.docx')


from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def create_book2(name, major, school, _time):
    # 生成一个文档
    doc1 = Document()
    # 增加标题
    title = doc1.add_paragraph()
    run = title.add_run('录取通知书')
    # 设置标题的样式
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.name = ''
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 将段落格式中的对齐方式设置为居中
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 增加内容
    doc1.add_paragraph(f'__{name}__ 同学：')
    content1 = doc1.add_paragraph(
        f'兹录取你入我校 __{major}__ 专业类学习。请凭本通知书来报道。具体时间、地点见《新生入学须知》。')
    # 设置内容样式
    content1.paragraph_format.first_line_indent = Pt(30)

    # 落款
    footer = doc1.add_paragraph(f'{school}\n')
    footer.add_run(f'{_time}')
    # 右对齐
    footer.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 保存文档
    doc1.save(f'07_录取通知书_{name}.docx')


if __name__ == '__main__':
    # create_book()
    create_book2('吕布', '人工智能技术', '清华大学', '二0三0年八月十号')


from win32com.client import constants, gencache


def createPdf(wordPath, pdfPath):
    """
    word转pdf
    :param wordPath: word文件路径
    :param pdfPath:  生成pdf文件路径
    """
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordPath, ReadOnly=1)
    doc.ExportAsFixedFormat(pdfPath,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)


if __name__ == "__main__":
    # 路径填写绝对路径
    createPdf(
        r'07_录取通知书_吕布.docx',
        r'10_word转换成pdf.pdf'
    )

from docx import Document


def read_word():
    # 打开文档
    doc1 = Document('人工智能.docx')
    # 读取数据-段落
    for p in doc1.paragraphs:
        print(p.text)
    # 读取表格
    for t in doc1.tables:
        for row in t.rows:
            for c in row.cells:
                print(c.text, end=' ')
            print()


if __name__ == '__main__':
    read_word()

"""按照标题序号选择要保留的段落"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

def delete_element(element):
    """
    删除文档中的元素，无论是表格还是段落
    """
    elt = element._element
    elt.getparent().remove(elt)

def iter_block_items(parent):
    """
    生成器，用于迭代Word文档中的块级项目（段落和表格）
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def delete_content_between_headings(doc_path, start_heading_index, end_heading_index):
    doc = Document(doc_path)
    current_heading_index = 0
    preserve_content = False
    elements_to_delete = []

    # 遍历文档元素，标记不需要保留的元素
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph) and block.style.name.startswith('Heading'):
            current_heading_index += 1
            if current_heading_index == start_heading_index:
                preserve_content = True  # 开始保留内容
            elif current_heading_index == end_heading_index:
                preserve_content = False  # 结束保留内容

        if not preserve_content:
            elements_to_delete.append(block)

    # 删除标记的元素
    for element in elements_to_delete:
        delete_element(element)
    # 在文档最后添加一个空格，再保存文档

    doc.save('modified_document.docx')


# 调用函数

delete_content_between_headings('人工智能 - 副本.docx', 3, 4)


"""合并文档，图片合不了，段落样式改变"""

from docx import Document
import os


def merge_word_documents(folder_path, output_path):
    # 创建一个新的Word文档
    merged_document = Document()

    # 遍历文件夹中的每个Word文档
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)

            # 打开当前的Word文档
            current_document = Document(file_path)

            # 将当前文档的内容复制到合并文档中
            for element in current_document.element.body:
                merged_document.element.body.append(element)

    # 保存合并后的文档
    merged_document.save(output_path)


# 程序入口
if __name__ == "__main__":
    folder_path = "output"
    output_path = "merged_document.docx"
    merge_word_documents(folder_path, output_path)


"""用spire.doc库合并文档的两个方法"""

# 方法1_通过插入文件来合并Word文档：使用该方法合并Word文档时，插入的文档默认从新的一页开始显示。
from spire.doc import *
from spire.doc.common import *

# 创建Document对象
doc = Document()

# 加载一个 Word 文档
doc.LoadFromFile("模板.docx")

# 将需合并的Word文档的内容插入到当前文档中
doc.InsertTextFromFile("new_copy.docx", FileFormat.Auto)
doc.InsertTextFromFile("人工智能.docx", FileFormat.Auto)

# 保存生成文档
doc.SaveToFile("合并Word文档.docx")
doc.Close()

#  方法2_通过复制内容合并 Word 文档：使用该方法合并Word文档时，各文档将紧接上一个文档的末尾。
from spire.doc import *
from spire.doc.common import *

# 加载第一个Word文档
doc1 = Document()
doc1.LoadFromFile("模板.docx")

# 将其余要合并的Word文档添加到列表中
files = []
files.append("new_copy.docx")
files.append("人工智能.docx")

# 获取第一个文档的最后一个节
lastSection = doc1.Sections.get_Item(doc1.Sections.Count - 1)

# 遍历列表中的文档
for file in files:
    doc = Document()
    doc.LoadFromFile(file)

    # 遍历每个文档中的各个节
    for i in range(doc.Sections.Count):
        section = doc.Sections.get_Item(i)

        # 遍历各个节中的子对象
        for j in range(section.Body.ChildObjects.Count):
            obj = section.Body.ChildObjects.get_Item(j)

            # 将每个文档中的子对象添加到第一个文档的最后一个节中
            lastSection.Body.ChildObjects.Add(obj.Clone())

# 保存合并后的文档
doc1.SaveToFile("合并Word文档2.docx")
doc1.Close()
doc.Close()

from spire.doc import *
from spire.doc.common import *

# 创建一个Document对象
document = Document()

# 加载名为"示例文档.docx"的Word文档
document.LoadFromFile("通用扩初模板.docx")

# 遍历文档的每个章节（Section）
for i in range(document.Sections.Count):
    # 获取当前索引的章节（Section）
    section = document.Sections.get_Item(i)

    # 构造结果文件的文件名，包括路径和文件名，格式为："按分节符拆分/结果文件_编号.docx"
    result = "按分节符拆分/" + "结果文件_{0}.docx".format(i + 1)

    # 创建一个新的Word文档对象来存储当前部分的内容
    newWord = Document()

    # 将当前章节（Section）的克隆添加到新Word文档的Sections中
    newWord.Sections.Add(section.Clone())

    # 将新Word文档保存到指定的文件名
    newWord.SaveToFile(result)
    newWord.Close()

document.Dispose()

from spire.doc import *
from spire.doc.common import *

# 创建一个Document对象
document = Document()

# 加载名为"示例文档.docx"的Word文档
document.LoadFromFile("人工智能 - 副本.docx")

# 初始化新文档和章节索引
newWord = None
file_index = 0

# 遍历文档的每个段落
for section in document.Sections:
    for paragraph in section.Paragraphs:
        # 检查段落样式是否为"Heading"
        if paragraph.StyleName.startswith("Heading"):
            # 如果当前有一个打开的文档，先保存并关闭它
            if newWord is not None:
                result = f"按分节符拆分/结果文件_{file_index}.docx"
                newWord.SaveToFile(result)
                newWord.Close()

            # 创建一个新的文档并增加章节索引
            newWord = Document()
            file_index += 1

        # 如果有打开的文档，将当前段落添加到新文档中
        if newWord is not None:
            new_section = newWord.AddSection()
            new_section.Paragraphs.Add(paragraph.Clone())

# 保存并关闭最后一个文档
if newWord is not None:
    result = f"按分节符拆分/结果文件_{file_index}.docx"
    newWord.SaveToFile(result)
    newWord.Close()

document.Dispose()

from pathlib import Path
from docxcompose.composer import Composer
from docx import Document

# 获取所有要合并的.docx文件的路径
result = []
for uu in Path('output').rglob('*.docx'):
    result.append(uu.resolve())

print("Files to merge:", result)

# 确保找到的文件列表不为空
if not result:
    print("No documents found to merge.")
    exit()

# 设置主文档
filename_master = result[0]
master = Document(filename_master)
composer = Composer(master)

# 遍历余下的文件，并将它们追加到主文档
for file_path in result[1:]:
    doc_temp = Document(file_path)
    doc_temp.add_page_break()
    composer.append(doc_temp)

# 保存合并后的文档
composer.save("merged_document.docx")
print("Documents successfully merged into 'merged_document.docx'")

import win32com.client as win32
from win32com.client import constants
import os

# 打开word应用程序
word = win32.gencache.EnsureDispatch('Word.Application')
# 是否可视化
word.Visible = 0
# 源文件路径
file_path = '人工智能 - 副本.docx'
# 打开
doc = word.Documents.Open(file_path)
# 光标start的查找
# 赋值对象
search_range = doc.Content
# 查找内容
search_range.Find.Execute(FindText="标题一")
# 选中查找到的内容
search_range.Select()
# 光标左移
word.Selection.MoveLeft()
# 将光标位置赋予start
start = word.Selection.Start.numerator
print(start)

# 光标end的查找  同上
search_range = doc.Content
search_range.Find.Execute(FindText="标题二")
search_range.Select()
word.Selection.MoveLeft()
end = word.Selection.Start.numerator
print(end)

# 选取光标start到光标end的内容
doc.Range(start, end).Select()
# 复制
word.Selection.Copy()
# 粘贴的目标文件
doc_new = word.Documents.Open('B.docx')
# 粘贴
doc_new.Application.ActiveDocument.Range().Paste()
# 关闭两个文件
doc_new.Close()
doc.Close()


