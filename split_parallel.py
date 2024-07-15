"""描述：通过文档标题结构创建层级文件夹
              复制原文档到每个子文件夹
              通过文件储存路径检索到文档中标题位置
              删除除这个标题下章节的外的所有内容'
              实现文档按标题拆分"""
import multiprocessing
import os
from shutil import copyfile
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import concurrent.futures
import time

def sanitize_filename(filename):
    """Clean up filename by replacing invalid characters with underscores."""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename

def create_folder_structure(doc_path, base_path="output"):
    doc = Document(doc_path)
    current_path = []
    paths = []

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            heading_text = sanitize_filename(para.text.strip())
            current_path = current_path[:level - 1] + [heading_text]
            path = os.path.join(base_path, *current_path)
            if not os.path.exists(path):
                os.makedirs(path)
            paths.append((path, para.text))

    return paths

#复制文件到没有子目录的文件夹
def copy_document_to_folders(doc_path, paths):
    for path, _ in paths:
        # Check if the directory has no subdirectories
        if not any(d for d in os.listdir(path) if os.path.isdir(os.path.join(path, d))):
            copyfile(doc_path, os.path.join(path, "document.docx"))

def delete_element(element):
    """
    删除文档中的元素，无论是表格还是段落
    """
    elt = element._element
    elt.getparent().remove(elt)

def iter_block_items(parent):
    """
    生成器，用于迭代Word文档中的块级项目（段落和表格）
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def delete_content_between_headings(doc_path, start_heading_index, end_heading_index):
    doc = Document(doc_path)
    current_heading_index = 0
    preserve_content = False
    elements_to_delete = []

    # 遍历文档元素，标记不需要保留的元素
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph) and block.style.name.startswith('Heading'):
            current_heading_index += 1
            if current_heading_index == start_heading_index:
                preserve_content = True  # 开始保留内容
            elif current_heading_index == end_heading_index:
                preserve_content = False  # 结束保留内容

        if not preserve_content:
            elements_to_delete.append(block)

    # 删除标记的元素


    for element in elements_to_delete:
        delete_element(element)


    #建立并行计算，在elements_to_delete中的元素运行delete_element
    #with concurrent.futures.ThreadPoolExecutor(max_workers=32) as executor:
    #    futures = [executor.submit(delete_element, element) for element in elements_to_delete]
    #    concurrent.futures.wait(futures)




    #删除文档里的段落样式为标题的内容
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            delete_element(para)




    doc.save(doc_path)



def process_document(file_path):
    # 加载Word文档
    doc = Document(file_path)

    # 提取文档中所有标题的文本
    headings = [para.text for para in doc.paragraphs if 'Heading' in para.style.name]

    # 找到start_heading在文档中的索引
    path_parts = file_path.split(os.sep)
    start_heading = path_parts[-2]  # 假设目录名称即为start_heading
    start_heading_index = headings.index(start_heading)

    # 初始化下一个标题的索引为start_heading_index之后的索引
    next_heading_index = start_heading_index + 1

    delete_content_between_headings(file_path, start_heading_index+1, next_heading_index+1)

def process_documents(base_path):
    # 收集所有需要处理的文档路径
    file_paths = []
    for root, dirs, files in os.walk(base_path):
        for file in files:
            if file.endswith(".docx") and file == "document.docx":
                file_path = os.path.join(root, file)
                file_paths.append(file_path)

    # 使用多进程处理文档
    with multiprocessing.Pool() as pool:
        pool.map(process_document, file_paths)



if __name__ == "__main__":
    tic = time.time()
    original_doc_path = "通用扩初模板.docx"
    paths = create_folder_structure(original_doc_path)
    copy_document_to_folders(original_doc_path, paths)
    process_documents("output")
    toc = time.time()
    print(f"Total time taken: {toc - tic:.2f} seconds.")